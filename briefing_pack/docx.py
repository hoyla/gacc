"""Docx output for the briefing pack — parallel surface to `03_Findings.md`.

Lisa-facing surface that carries charts; the markdown stays canonical
(NotebookLM-feed, per `memory/architecture_journalist_surfaces.md` —
keep LLM output / images / interpretation OUTSIDE documents downstream
LLM tools will read). Verified Drive → Google Docs round-trip fidelity
2026-05-16; see `dev_notes/2026-05-16_docx-drive-spike.md`.

v1 (this slice): top-N movers as a stack of cards — one heading +
paragraph + 24-month line chart per mover. Chart shows the prior-12mo
window in grey and the current-12mo window in red, sharing an x-axis
of 24 calendar months. Data source: `eurostat_raw_rows` filtered to
the finding's `(flow, partners, hs_patterns, scope=eu_27)`.

Caller contract: `render_top_movers_docx(out_path)` opens its own DB
connection, fetches movers via the same `_compute_top_movers` helper
the markdown renderer uses, and writes the .docx atomically.
"""

from __future__ import annotations

import io
import logging
import math
from datetime import date, datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # non-interactive backend; deterministic on a
                       # given host. See dev_notes design doc step 5
                       # for the known cross-host determinism caveat.
import matplotlib.pyplot as plt
import psycopg2.extras
from docx import Document
from docx.shared import Mm, Pt

from briefing_pack._helpers import (
    DEFAULT_TOP_N,
    _compute_predictability_per_group,
    _compute_top_movers,
    _conn,
    _fmt_eur,
    _fmt_pct,
)

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Page-setup constants — values verified by the 2026-05-16 docx fidelity
# spike to round-trip cleanly through Drive → Google Docs.
# ---------------------------------------------------------------------------

_PAGE_WIDTH_MM = 210     # A4 portrait
_PAGE_HEIGHT_MM = 297    # A4 portrait
_MARGIN_MM = 10          # all four sides
_BODY_FONT_PT = 11
_CHART_WIDTH_MM = 190    # 210 - 2×10 usable width, no breathing margin

# EU-27 reporter exclusion (matches anomalies.EU27_EXCLUDE_REPORTERS).
# UK left the EU 2020-01; eurostat_raw_rows still includes GB rows for
# UK-side enquiries, but eu_27-scope findings exclude them. Hardcoded
# here rather than imported from anomalies to keep briefing_pack from
# growing an analyser-module dependency.
_EU27_EXCLUDE_REPORTERS = ("GB",)


# ---------------------------------------------------------------------------
# Date helpers
# ---------------------------------------------------------------------------

def _months_back(d: date, n: int) -> date:
    """Return `date(d - n months, day=1)` without pulling in dateutil.

    Year/month arithmetic via month-index. 23 months back from
    2026-02-01 → 2024-03-01.
    """
    total = d.year * 12 + (d.month - 1) - n
    return date(total // 12, (total % 12) + 1, 1)


_MONTH_LABEL = {
    1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
    7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec",
}


def _month_iter(start: date, end: date):
    """Yield consecutive month-anchored dates from `start` to `end` inclusive."""
    cur = date(start.year, start.month, 1)
    end_anchor = date(end.year, end.month, 1)
    while cur <= end_anchor:
        yield cur
        # next month
        if cur.month == 12:
            cur = date(cur.year + 1, 1, 1)
        else:
            cur = date(cur.year, cur.month + 1, 1)


# ---------------------------------------------------------------------------
# Data fetch — finding detail + monthly chart series
# ---------------------------------------------------------------------------

def _fetch_finding_detail(cur, finding_id: int) -> dict | None:
    """Load the full `detail` JSONB for a finding by id.

    `_compute_top_movers` returns only the top-line columns; for the
    chart we need `method_query.flow`, `.partners`, `.hs_patterns`.
    Cheap re-fetch via the primary key.
    """
    cur.execute(
        "SELECT detail FROM findings WHERE id = %s",
        (finding_id,),
    )
    row = cur.fetchone()
    if row is None:
        return None
    return row[0] if not isinstance(row, dict) else row["detail"]


def _fetch_monthly_eur_series(
    cur,
    *,
    hs_patterns: list[str],
    flow: int,
    partners: list[str],
    start: date,
    end: date,
) -> dict[date, float]:
    """Return {month_date: value_eur} summing `eurostat_raw_rows` over the
    window. Months with no data are absent from the dict (not zeroed) so
    the caller can decide gap-rendering policy.

    Matches the filter shape of `_hs_group_top_cn8s` in `anomalies.py`:
    flow ∈ {1=import, 2=export}, partners filter, HS pattern OR-LIKE,
    and exclude `_EU27_EXCLUDE_REPORTERS` from the reporter set.
    """
    if not hs_patterns:
        return {}
    # OR'd LIKE matches the existing analyser pattern (see
    # `anomalies._hs_pattern_or_clause` — the index-friendly form;
    # avoid LIKE ANY which forces a seq scan).
    like_clause = "(" + " OR ".join(
        ["product_nc LIKE %s"] * len(hs_patterns)
    ) + ")"
    sql = f"""
        SELECT date_trunc('month', period)::date AS month,
               SUM(value_eur)::float8 AS value_eur
          FROM eurostat_raw_rows
         WHERE period >= %s AND period <= %s
           AND flow = %s
           AND partner = ANY(%s)
           AND {like_clause}
           AND reporter <> ALL(%s)
      GROUP BY 1
      ORDER BY 1
    """
    params = (
        start, end, flow, list(partners),
        *hs_patterns,
        list(_EU27_EXCLUDE_REPORTERS),
    )
    cur.execute(sql, params)
    return {r[0]: float(r[1] or 0.0) for r in cur.fetchall()}


# ---------------------------------------------------------------------------
# Chart rendering
# ---------------------------------------------------------------------------

def _pick_eur_scale(max_value: float) -> tuple[float, str]:
    """Pick a divisor + label for a y-axis based on the series' max.

    Avoids matplotlib's default `1e9` exponent annotation by scaling
    explicitly to billions / millions / thousands.
    """
    if max_value >= 1e9:
        return 1e9, "€ billions"
    if max_value >= 1e6:
        return 1e6, "€ millions"
    if max_value >= 1e3:
        return 1e3, "€ thousands"
    return 1.0, "€"


def _build_chart_png(
    *,
    current_end: date,
    monthly_eur: dict[date, float],
    group_name: str,
    flow_label: str,
) -> bytes:
    """Render a 24-month line chart, prior-12mo grey vs current-12mo red.

    `monthly_eur` is a {month → value} dict; missing months render as
    gaps in the line (matplotlib handles NaN by skipping).

    Returns PNG bytes ready to drop into `doc.add_picture()`.
    """
    # Build the 24-month axis ending at current_end, inclusive.
    start = _months_back(current_end, 23)
    months = list(_month_iter(start, current_end))
    values = [monthly_eur.get(m, float("nan")) for m in months]

    # Split into prior 12mo (first 12 points) + current 12mo (last 12).
    # Boundary month is the 13th — the last month of the prior window.
    prior_vals = values[:12]
    current_vals = [float("nan")] * 11 + values[11:]  # overlap on
                                                       # the boundary
                                                       # so both lines
                                                       # are visible
    labels = [
        f"{_MONTH_LABEL[m.month]} {m.year % 100:02d}" for m in months
    ]

    scale, unit_label = _pick_eur_scale(
        max((v for v in values if not math.isnan(v)), default=0.0)
    )
    prior_scaled = [v / scale if not math.isnan(v) else float("nan")
                    for v in prior_vals]
    current_scaled = [v / scale if not math.isnan(v) else float("nan")
                      for v in current_vals]

    fig, ax = plt.subplots(figsize=(7.5, 3.2), dpi=150)
    ax.plot(
        labels[:12], prior_scaled,
        label="Prior 12mo", linewidth=2, color="#888888",
    )
    ax.plot(
        labels, current_scaled,
        label="Current 12mo", linewidth=2.5, color="#cc3333",
    )
    title = f"{group_name} — {flow_label}"
    ax.set_title(title, fontsize=11, loc="left")
    ax.set_ylabel(unit_label, fontsize=9)
    ax.tick_params(axis="x", labelsize=7, rotation=45)
    ax.tick_params(axis="y", labelsize=8)
    ax.grid(True, alpha=0.3)
    ax.legend(loc="best", fontsize=8, frameon=False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Docx assembly
# ---------------------------------------------------------------------------

def _apply_page_setup(doc: Document) -> None:
    """Apply A4 + 10mm margins + 11pt body to a fresh Document."""
    section = doc.sections[0]
    section.page_height = Mm(_PAGE_HEIGHT_MM)
    section.page_width = Mm(_PAGE_WIDTH_MM)
    section.top_margin = Mm(_MARGIN_MM)
    section.bottom_margin = Mm(_MARGIN_MM)
    section.left_margin = Mm(_MARGIN_MM)
    section.right_margin = Mm(_MARGIN_MM)
    doc.styles["Normal"].font.size = Pt(_BODY_FONT_PT)


def _flow_label_for_subkind(subkind: str) -> str:
    """Human-readable flow label matching the markdown renderer's convention."""
    if subkind.endswith("_export"):
        return "EU-27 exports (reporter→CN)"
    return "EU-27 imports (CN→reporter)"


def render_top_movers_docx(
    out_path: str | Path,
    *,
    top_n: int = DEFAULT_TOP_N,
    scope_label: str | None = None,
) -> Path:
    """Render the top-N movers section of the briefing pack to a docx.

    Each mover gets a card: H2 heading with the group name + badge, a
    paragraph with the headline figures, and a 24-month line chart
    showing prior-12mo vs current-12mo monthly EUR series. Chart data
    is fetched from `eurostat_raw_rows` using the same filter the
    analyser used to produce the finding (flow, partners, HS patterns,
    EU-27 reporter set).

    Returns the resolved Path of the written file.

    Empty-movers case (no findings pass the editorial filter): writes a
    title + single italic paragraph rather than an empty doc, so the
    bundle artefact is always valid.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with _conn() as conn, conn.cursor(
        cursor_factory=psycopg2.extras.DictCursor,
    ) as cur:
        predictability = _compute_predictability_per_group(cur)
        movers = _compute_top_movers(
            cur, predictability=predictability, limit=top_n,
        )

        doc = Document()
        _apply_page_setup(doc)

        title_text = "Meridian — Findings"
        if scope_label:
            title_text += f" ({scope_label})"
        doc.add_heading(title_text, level=0)

        p = doc.add_paragraph()
        p.add_run("Generated: ").bold = True
        p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

        if not movers:
            empty_p = doc.add_paragraph()
            empty_p.add_run(
                "No top-mover findings passed the editorial filter for "
                "this cycle. See 03_Findings.md (Tier 2) for the full "
                "state of play."
            ).italic = True
            doc.save(str(out_path))
            log.info("Wrote findings docx to %s (no movers)", out_path)
            return out_path

        doc.add_heading(f"Top {len(movers)} movers this cycle", level=1)

        preamble = doc.add_paragraph()
        preamble.add_run(
            "Editorially-quotable shifts ranked by a composite of "
            "|YoY| × log(€). Filters: ≥10pp move, ≥€100M current 12mo "
            "total, not low-base, predictability badge ≠ 🔴. Each "
            "chart shows the prior 12-month window (grey) and the "
            "current 12-month window (red) summed monthly. Figures "
            "match 03_Findings.md."
        ).italic = True

        # Re-use this cursor for the per-finding chart data fetches.
        for m in movers:
            flow_label = _flow_label_for_subkind(m["subkind"])
            pred = m.get("predictability")
            badge = f" {pred[0]}" if pred is not None else ""
            yoy_kg = m.get("yoy_pct_kg")
            kg_str = (
                f" (kg {_fmt_pct(yoy_kg)})" if yoy_kg is not None else ""
            )
            period = m["current_end"]

            doc.add_heading(f"{m['group_name']}{badge}", level=2)

            card_p = doc.add_paragraph()
            card_p.add_run(f"{flow_label}: ").bold = True
            card_p.add_run(
                f"{_fmt_pct(m['yoy_pct'])}{kg_str} to "
                f"{_fmt_eur(m['current_eur'])} "
                f"(12mo to {period.strftime('%Y-%m')}). "
            )
            card_p.add_run(f"finding/{m['id']}").italic = True

            # Chart: re-fetch detail to get hs_patterns + method_query.
            detail = _fetch_finding_detail(cur, m["id"])
            chart_inserted = False
            if detail:
                method_q = detail.get("method_query", {})
                hs_patterns = (
                    method_q.get("hs_patterns")
                    or detail.get("group", {}).get("hs_patterns")
                    or []
                )
                flow = int(method_q.get("flow") or 0)
                partners = method_q.get("partners") or []
                start = _months_back(period, 23)
                if hs_patterns and flow and partners:
                    series = _fetch_monthly_eur_series(
                        cur,
                        hs_patterns=hs_patterns,
                        flow=flow,
                        partners=partners,
                        start=start,
                        end=period,
                    )
                    if series:
                        png = _build_chart_png(
                            current_end=period,
                            monthly_eur=series,
                            group_name=m["group_name"],
                            flow_label=flow_label,
                        )
                        doc.add_picture(
                            io.BytesIO(png), width=Mm(_CHART_WIDTH_MM),
                        )
                        chart_inserted = True

            if not chart_inserted:
                miss = doc.add_paragraph()
                miss.add_run(
                    "(Chart unavailable — finding detail or "
                    "underlying observations not fetchable.)"
                ).italic = True

    doc.save(str(out_path))
    log.info(
        "Wrote findings docx to %s (%d movers with charts)",
        out_path,
        len(movers),
    )
    return out_path
